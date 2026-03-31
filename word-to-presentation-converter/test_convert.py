from docx import Document
from app import parse_docx, create_pptx, _is_probable_title_block, _normalize_extracted_text

# Regression checks: don't treat the first numbered question as a title
assert _is_probable_title_block('Maths Quiz Title') is True
assert _is_probable_title_block('1. What is 457 + 29 + 158 = ________ ?') is False
assert _is_probable_title_block('A) 624 .') is False

# Regression: preserve mixed fractions commonly mangled by PDF text extraction
assert _normalize_extracted_text('What is 1 1 2 + 2 3 4 ?') == 'What is 1 1/2 + 2 3/4 ?'
assert _normalize_extracted_text('A) 3 5 6 .') == 'A) 3 5/6 .'
assert _normalize_extracted_text('B) 11 6') == 'B) 11/6'
assert _normalize_extracted_text('Answer: A) 34 35') == 'Answer: A) 34/35'

# create sample docx
doc = Document()
doc.add_paragraph('Maths Quiz Title')
doc.add_paragraph('')
# question with options and 'a.)' style labels and extra blank lines
doc.add_paragraph('1. What is 2+2?')
doc.add_paragraph('')
doc.add_paragraph('a.) 3')
doc.add_paragraph('')
doc.add_paragraph('b.) 4')
doc.add_paragraph('')
doc.add_paragraph('c.) 5')
doc.add_paragraph('')
doc.add_paragraph('d.) 6')

# long question that would wrap
long_q = 'Solve the following equation: ' + ('x+'*200)
doc.add_paragraph('')
doc.add_paragraph(long_q)

doc.save('sample_test.docx')

# parse and convert
with open('sample_test.docx','rb') as f:
    title, qs = parse_docx(f)
    pptx_io = create_pptx(title, qs)
    with open('sample_output.pptx','wb') as out:
        out.write(pptx_io.read())

print('Wrote sample_output.pptx')
